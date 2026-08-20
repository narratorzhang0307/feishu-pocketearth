#!/usr/bin/env python3
"""Build the marked-up Pocket Earth Google-edition full transcript DOCX."""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/zhangcheng/Desktop/pocket earth_google")
VOICE_DIR = ROOT / "PocketEarthGoogle_提交包" / "03_视频口播"
FINAL_SRT = VOICE_DIR / "02_Google版_全片字幕.srt"
ORIGINAL_SRT = VOICE_DIR / "01_原片完整字幕_人工校订版.srt"
OUTPUT_DOCX = VOICE_DIR / "Google版全片字幕_红色修订标记.docx"


# Selected design system: compact_reference_guide.
# Selected first-page pattern: memo_masthead, adapted as a subtitle working draft.
# Named overrides are documented where they are introduced below.
# Named compatibility override: use a macOS system CJK family for all script ranges
# so both Microsoft Word and the LibreOffice QA renderer preserve Simplified Chinese.
LATIN_FONT = "Hiragino Sans GB"
CJK_FONT = "Hiragino Sans GB"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x18, 0x2B, 0x3A)
BLACK = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x66, 0x6B, 0x73)
LIGHT_MUTED = RGBColor(0x8A, 0x90, 0x99)
REVISION_RED = RGBColor(0xC6, 0x28, 0x28)
PALE_RED = "FDECEC"
PALE_BLUE = "E8EEF5"


@dataclass(frozen=True)
class Cue:
    index: int
    start_ms: int
    end_ms: int
    start: str
    end: str
    text: str


@dataclass(frozen=True)
class ChangeBlock:
    start_ms: int
    end_ms: int
    title: str
    purpose: str


CHANGE_BLOCKS = [
    ChangeBlock(149_080, 188_260, "技术总览与 Gemini 云端分工", "将旧模型网关主线改为 Google-first 推理路线。"),
    ChangeBlock(224_580, 253_020, "跨文化看展闭环", "明确 Gemini 跨文化理解、证据标注与确认后写入。"),
    ChangeBlock(260_920, 281_700, "Agent Forge 真实性口径", "删除未实现的联网研究表述，改为候选、核验和确认机制。"),
    ChangeBlock(296_240, 368_780, "Gemma 3n E2B IT + MediaPipe 端侧推理", "写清 int4 Web 权重、.litertlm、WebGPU、真实加载验证与 LiteRT-LM 迁移边界。"),
    ChangeBlock(389_820, 423_180, "Google-first 总结与出海价值", "用 Gemma、Gemini 与跨文化同理心收束全片。"),
]


def parse_timestamp(value: str) -> int:
    hours, minutes, rest = value.split(":")
    seconds, millis = rest.split(",")
    return (
        int(hours) * 3_600_000
        + int(minutes) * 60_000
        + int(seconds) * 1_000
        + int(millis)
    )


def parse_srt(path: Path) -> list[Cue]:
    raw = path.read_text(encoding="utf-8-sig").strip()
    chunks = re.split(r"\n\s*\n", raw)
    cues: list[Cue] = []
    for chunk in chunks:
        lines = [line.rstrip() for line in chunk.splitlines()]
        if len(lines) < 3 or "-->" not in lines[1]:
            continue
        start, end = [item.strip() for item in lines[1].split("-->", 1)]
        cues.append(
            Cue(
                index=int(lines[0]),
                start_ms=parse_timestamp(start),
                end_ms=parse_timestamp(end),
                start=start,
                end=end,
                text=" ".join(line.strip() for line in lines[2:] if line.strip()),
            )
        )
    return cues


def overlaps(cue: Cue, block: ChangeBlock) -> bool:
    return cue.start_ms < block.end_ms and cue.end_ms > block.start_ms


def is_changed(cue: Cue) -> bool:
    return any(overlaps(cue, block) for block in CHANGE_BLOCKS)


def format_short_time(ms: int) -> str:
    total_seconds = ms // 1_000
    minutes, seconds = divmod(total_seconds, 60)
    return f"{minutes:02d}:{seconds:02d}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    strike: bool | None = None,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if strike is not None:
        run.font.strike = strike


def set_paragraph_spacing(paragraph, before: float, after: float, line: float) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # Named override: compact subtitle cues need tighter spacing than ordinary body prose.
    cue_style = doc.styles.add_style("Transcript Cue", WD_STYLE_TYPE.PARAGRAPH)
    cue_style.base_style = normal
    cue_style.font.name = LATIN_FONT
    cue_style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    cue_style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    cue_style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    cue_style.font.size = Pt(10.5)
    cue_style.paragraph_format.space_before = Pt(0)
    cue_style.paragraph_format.space_after = Pt(4)
    cue_style.paragraph_format.line_spacing = 1.15
    cue_style.paragraph_format.keep_together = True

    # Named override: source/target revision paragraphs carry long spoken text.
    revision_style = doc.styles.add_style("Revision Body", WD_STYLE_TYPE.PARAGRAPH)
    revision_style.base_style = normal
    revision_style.font.name = LATIN_FONT
    revision_style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    revision_style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    revision_style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    revision_style.font.size = Pt(10.5)
    revision_style.paragraph_format.space_before = Pt(0)
    revision_style.paragraph_format.space_after = Pt(9)
    revision_style.paragraph_format.line_spacing = 1.25


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=LIGHT_MUTED)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(hp, 0, 0, 1.0)
    run = hp.add_run("POCKET EARTH  ·  GOOGLE 提交版视频口播")
    set_run_font(run, size=8.5, color=LIGHT_MUTED, bold=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(fp, 0, 0, 1.0)
    run = fp.add_run("Google 版全片字幕  |  第 ")
    set_run_font(run, size=8.5, color=LIGHT_MUTED)
    add_page_field(fp)
    run = fp.add_run(" 页")
    set_run_font(run, size=8.5, color=LIGHT_MUTED)


def add_title_block(doc: Document, changed_count: int) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 8, 3, 1.0)
    r = p.add_run("Pocket Earth · Google 版")
    set_run_font(r, size=25, color=NAVY, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 14, 1.0)
    r = p.add_run("全片字幕（红色修订标记）")
    set_run_font(r, size=14, color=MUTED, bold=False)

    metadata = [
        ("用途", "最终剪辑、重录口播与技术审核核对"),
        ("成片时长", "约 07:23"),
        ("修订范围", f"5 个替换区间，共 {changed_count} 条 Google 版红色字幕"),
        ("版本日期", date(2026, 7, 14).strftime("%Y 年 %-m 月 %-d 日")),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 2, 1.1)
        r = p.add_run(f"{label}：")
        set_run_font(r, size=10.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 13, 4, 1.15)
    r = p.add_run("标记规则  ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run("黑色 = 沿用原片口播；")
    set_run_font(r, size=10.5, color=BLACK)
    r = p.add_run("红色 = Google 版新增或替换；")
    set_run_font(r, size=10.5, color=REVISION_RED, bold=True)
    r = p.add_run("红色删除线 = 原稿中需删除的内容。")
    set_run_font(r, size=10.5, color=REVISION_RED, strike=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 8, 2, 1.1)
    set_paragraph_shading(p, PALE_RED)
    r = p.add_run("技术审核重点  ")
    set_run_font(r, size=10.5, color=REVISION_RED, bold=True)
    r = p.add_run("Google Gemma 3n E2B IT · int4 Web（.litertlm）已安装到项目；MediaPipe LLM Inference Web 通过 WebGPU 完成浏览器端推理，并已实测加载与端侧生成。")
    set_run_font(r, size=10.5, color=REVISION_RED, bold=True)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 7, 1.15)
    r = p.add_run("提示：正文是最终 Google 版全片字幕；附录保留原稿删改对照，删除线内容不进入最终成片。")
    set_run_font(r, size=9.5, color=MUTED, italic=True)


def add_full_transcript(doc: Document, cues: list[Cue]) -> None:
    h = doc.add_paragraph("一、Google 版全片字幕", style="Heading 1")
    h.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 8, 1.15)
    r = p.add_run("按时间轴顺序使用。红色句子为重录/替换内容，其余黑色句子沿用原片。")
    set_run_font(r, size=9.5, color=MUTED)

    for cue in cues:
        changed = is_changed(cue)
        p = doc.add_paragraph(style="Transcript Cue")
        r = p.add_run(f"{cue.start} – {cue.end}   ")
        set_run_font(r, size=8.5, color=LIGHT_MUTED, bold=True)
        r = p.add_run(cue.text)
        set_run_font(
            r,
            size=10.5,
            color=REVISION_RED if changed else BLACK,
            bold=changed,
        )


def gather_block_text(cues: list[Cue], block: ChangeBlock) -> str:
    return " ".join(cue.text for cue in cues if overlaps(cue, block))


def add_revision_appendix(
    doc: Document,
    original_cues: list[Cue],
    final_cues: list[Cue],
) -> None:
    doc.add_page_break()
    h = doc.add_paragraph("二、删改对照（五个替换区间）", style="Heading 1")
    h.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 10, 1.15)
    r = p.add_run("每段先列原稿删除内容，再列 Google 版最终口播；两者均用红色标出。")
    set_run_font(r, size=9.5, color=MUTED)

    for idx, block in enumerate(CHANGE_BLOCKS, start=1):
        if idx > 1:
            doc.add_page_break()

        time_range = f"{format_short_time(block.start_ms)}–{format_short_time(block.end_ms)}"
        h = doc.add_paragraph(f"改动 {idx}｜{time_range}｜{block.title}", style="Heading 2")
        h.paragraph_format.keep_with_next = True

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 8, 1.15)
        r = p.add_run("修改目的：")
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        r = p.add_run(block.purpose)
        set_run_font(r, size=9.5, color=MUTED)

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 3, 3, 1.0)
        r = p.add_run("原稿删除 / 替换")
        set_run_font(r, size=10.5, color=REVISION_RED, bold=True)
        p.paragraph_format.keep_with_next = True

        p = doc.add_paragraph(style="Revision Body")
        r = p.add_run(gather_block_text(original_cues, block))
        set_run_font(r, size=10.5, color=REVISION_RED, strike=True)

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 3, 3, 1.0)
        r = p.add_run("Google 版新增 / 替换")
        set_run_font(r, size=10.5, color=REVISION_RED, bold=True)
        p.paragraph_format.keep_with_next = True

        p = doc.add_paragraph(style="Revision Body")
        r = p.add_run(gather_block_text(final_cues, block))
        set_run_font(r, size=10.5, color=REVISION_RED, bold=True)


def save_document() -> Path:
    final_cues = parse_srt(FINAL_SRT)
    original_cues = parse_srt(ORIGINAL_SRT)
    if len(final_cues) != 125:
        raise ValueError(f"Expected 125 final cues, found {len(final_cues)}")
    if not original_cues:
        raise ValueError("Original subtitle file produced no cues")

    changed_count = sum(1 for cue in final_cues if is_changed(cue))
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_title_block(doc, changed_count)
    add_full_transcript(doc, final_cues)
    add_revision_appendix(doc, original_cues, final_cues)

    props = doc.core_properties
    props.title = "Pocket Earth · Google 版全片字幕（红色修订标记）"
    props.subject = "Google 提交版视频口播与删改对照"
    props.author = "Pocket Earth"
    props.keywords = "Pocket Earth, Google, Gemini, Gemma, MediaPipe, 字幕, 口播"

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    try:
        output = save_document()
        print(output)
    except Exception as exc:  # pragma: no cover - command-line diagnostics
        print(f"error: {exc}", file=sys.stderr)
        raise
