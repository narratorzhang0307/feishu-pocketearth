#!/usr/bin/env python3
"""Build the styled DOCX companion for the Photos aesthetic-curator execution plan."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "docs/strategy/Photos-审美选片LoRA与个人偏好学习执行计划.md"
OUTPUT = ROOT / "docs/strategy/Photos-审美选片LoRA与个人偏好学习执行计划.docx"
DOCUMENT_HEADER = "POCKET EARTH  /  PHOTOS · AESTHETIC CURATOR"
DOCUMENT_KICKER = "POCKET EARTH · RESEARCH / PRODUCT / DEVICE DELIVERY"
DOCUMENT_SUBTITLE = "专项执行计划 · Qwen3-VL-2B / Visual LoRA / MNN / 本地偏好学习"
CORE_TITLE = "Pocket Earth Photos 审美选片 LoRA 与个人偏好学习执行计划"
CORE_SUBJECT = "Qwen3-VL-2B、Visual LoRA、MNN、SME2 与本地个人偏好学习"
CORE_KEYWORDS = "Pocket Earth, Photos, Qwen3-VL, LoRA, MNN, SME2, 审美选片, 个人偏好"
CORE_COMMENTS = "由 Markdown 权威稿生成；正式实施前需修订总计划中的旧 LoRA 约束。"

PAGE_WIDTH_DXA = 9360
BODY_FONT = "Aptos"
CJK_FONT = "PingFang SC"
CODE_FONT = "Menlo"

INK = "171A1D"
MUTED = "687078"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
GREEN = "00B96B"
GREEN_LIGHT = "E8F8EF"
GRAY_LIGHT = "F3F5F6"
GRAY_BORDER = "C8CED3"
GOLD_LIGHT = "FFF4D6"
RED_LIGHT = "FDECEC"
WHITE = "FFFFFF"


def set_run_font(run, name: str = BODY_FONT, size: float | None = None, color: str | None = None) -> None:
    run.font.name = name
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), CJK_FONT if name == BODY_FONT else name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"table width must be {PAGE_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_shading(run, fill: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    shd = r_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        r_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 18, space: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_hyperlink(paragraph, text: str, url: str):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")


def add_inline(paragraph, text: str, base_size: float = 10.5, base_color: str = INK) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_run_font(run, size=base_size, color=base_color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, CODE_FONT, max(8.0, base_size - 1.0), DEEP_BLUE)
            set_run_shading(run, "EDF1F4")
        elif part.startswith("["):
            match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", part)
            if match:
                add_hyperlink(paragraph, match.group(1), match.group(2))
            else:
                run = paragraph.add_run(part)
                set_run_font(run, size=base_size, color=base_color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size, color=base_color)


def new_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"AESTH{abstract_id:03X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "number" else "☐" if kind == "check" else "•")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "480")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")
    ind.set(qn("w:hanging"), "240")
    p_pr.extend([tabs, ind])
    lvl.extend([start, fmt, lvl_text, suff, lvl_jc, p_pr])
    if kind != "number":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DEEP_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.styles["Title"]
    title.font.name = BODY_FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles.add_style("Plan Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = BODY_FONT
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    code = doc.styles.add_style("Plan Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = CODE_FONT
    code._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT)
    code.font.size = Pt(8.2)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = True


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(DOCUMENT_HEADER)
    run.bold = True
    set_run_font(run, CODE_FONT, 7.5, MUTED)
    set_paragraph_border(p, "bottom", GREEN, size=8, space=3)

    footer = section.footer
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("v1.0 · 2026-08-11    |    ")
    set_run_font(run, CODE_FONT, 7.5, MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), CODE_FONT)
    r_fonts.set(qn("w:hAnsi"), CODE_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.extend([r_fonts, color])
    field_run.extend([r_pr, fld_begin, instr, fld_sep, fld_text, fld_end])
    p._p.append(field_run)


def add_masthead(doc: Document, title: str, metadata: list[str]) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run(DOCUMENT_KICKER)
    run.bold = True
    set_run_font(run, CODE_FONT, 8.5, GREEN)

    title_p = doc.add_paragraph(title, style="Title")
    set_paragraph_border(title_p, "bottom", INK, size=18, space=7)

    subtitle = doc.add_paragraph(DOCUMENT_SUBTITLE, style="Plan Subtitle")
    subtitle.paragraph_format.keep_with_next = True

    pairs: list[tuple[str, str]] = []
    for line in metadata:
        clean = line.strip().rstrip("  ")
        clean = clean[1:].strip() if clean.startswith(">") else clean
        if "：" in clean:
            key, value = clean.split("：", 1)
            pairs.append((key.strip(), value.strip().replace("`", "")))
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [1750, 7610])
    for i, (key, value) in enumerate(pairs):
        left, right = table.rows[i].cells
        left.text = ""
        right.text = ""
        shade_cell(left, INK if i == 0 else DEEP_BLUE)
        shade_cell(right, GOLD_LIGHT if key == "规范关系" else GRAY_LIGHT)
        lp = left.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(key.upper())
        lr.bold = True
        set_run_font(lr, CODE_FONT, 7.8, WHITE)
        rp = right.paragraphs[0]
        rp.paragraph_format.space_after = Pt(0)
        add_inline(rp, value, 8.5, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def display_width(text: str) -> int:
    clean = re.sub(r"[`*_]", "", text)
    return sum(2 if ord(ch) > 127 else 1 for ch in clean)


def choose_widths(rows: list[list[str]]) -> list[int]:
    count = len(rows[0])
    max_lengths = [max(display_width(row[i]) for row in rows) for i in range(count)]
    caps = {2: 46, 3: 34, 4: 27, 5: 22}
    mins = {2: 1800, 3: 1300, 4: 1050, 5: 900}
    cap = caps.get(count, 20)
    minimum = mins.get(count, 800)
    scores = [max(5, min(length, cap)) for length in max_lengths]
    remaining = PAGE_WIDTH_DXA - minimum * count
    total = sum(scores)
    widths = [minimum + int(remaining * score / total) for score in scores]
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, rows[0]):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_inline(p, text, 8.2, WHITE)
        for run in p.runs:
            run.bold = True
        shade_cell(cell, DEEP_BLUE)
    repeat_header(table.rows[0])
    for row_index, values in enumerate(rows[1:], 1):
        cells = table.add_row().cells
        for col_index, (cell, text) in enumerate(zip(cells, values)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if re.fullmatch(r"[+≥<P0-9.%—– /]+", text.strip()):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, text, 8.2, INK)
            if row_index % 2 == 0:
                shade_cell(cell, GRAY_LIGHT)
            if col_index == 0:
                for run in p.runs:
                    run.bold = True
    set_table_geometry(table, choose_widths(rows))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_quote(doc: Document, text: str, first: bool) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4 if first else 0)
    p.paragraph_format.space_after = Pt(7)
    set_paragraph_shading(p, GREEN_LIGHT)
    set_paragraph_border(p, "left", GREEN, size=28, space=8)
    add_inline(p, text, 10.4, INK)
    for run in p.runs:
        run.bold = True


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Plan Code")
    set_paragraph_shading(p, GRAY_LIGHT)
    set_paragraph_border(p, "left", DEEP_BLUE, size=18, space=6)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break(WD_BREAK.LINE)
        run = p.add_run(line)
        set_run_font(run, CODE_FONT, 8.2, INK)


def add_section_heading(doc: Document, level: int, text: str) -> None:
    style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
    p = doc.add_paragraph(style=style)
    add_inline(p, text, 16 if level == 2 else 13 if level == 3 else 11.5, BLUE if level < 4 else DEEP_BLUE)
    if level == 2:
        set_paragraph_border(p, "bottom", GRAY_BORDER, size=6, space=4)


def parse_markdown(doc: Document, lines: list[str]) -> None:
    i = 0
    number_id: int | None = None
    bullet_id: int | None = None
    check_id: int | None = None
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            number_id = bullet_id = check_id = None
            i += 1
            continue
        if stripped.startswith("```"):
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, code)
            i += 1
            number_id = bullet_id = check_id = None
            continue
        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_border(p, "bottom", GRAY_BORDER, size=6, space=2)
            number_id = bullet_id = check_id = None
            i += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            add_section_heading(doc, len(heading.group(1)), heading.group(2))
            number_id = bullet_id = check_id = None
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            raw_rows = [split_table_row(line) for line in table_lines]
            rows = [row for index, row in enumerate(raw_rows) if index != 1 or not all(re.fullmatch(r":?-{3,}:?", cell) for cell in row)]
            add_markdown_table(doc, rows)
            number_id = bullet_id = check_id = None
            continue
        if stripped.startswith(">"):
            add_quote(doc, stripped[1:].strip(), True)
            number_id = bullet_id = check_id = None
            i += 1
            continue
        checklist = re.match(r"^-\s+\[\s\]\s+(.+)$", stripped)
        if checklist:
            if check_id is None:
                check_id = new_numbering(doc, "check")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            apply_numbering(p, check_id)
            add_inline(p, checklist.group(1), 10.2, INK)
            number_id = bullet_id = None
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if number_id is None:
                number_id = new_numbering(doc, "number")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            apply_numbering(p, number_id)
            add_inline(p, numbered.group(1), 10.2, INK)
            bullet_id = check_id = None
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            if bullet_id is None:
                bullet_id = new_numbering(doc, "bullet")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            apply_numbering(p, bullet_id)
            add_inline(p, bullet.group(1), 10.2, INK)
            number_id = check_id = None
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline(p, stripped, 10.5, INK)
        if stripped.endswith(("：", ":")):
            p.paragraph_format.keep_with_next = True
        number_id = bullet_id = check_id = None
        i += 1


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = CORE_TITLE
    props.subject = CORE_SUBJECT
    props.author = "Pocket Earth 决赛项目组"
    props.keywords = CORE_KEYWORDS
    props.comments = CORE_COMMENTS


def build() -> None:
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise RuntimeError("Markdown 缺少一级标题")
    title = lines[0][2:].strip()
    metadata: list[str] = []
    index = 1
    while index < len(lines):
        if lines[index].strip().startswith(">"):
            metadata.append(lines[index])
        elif lines[index].strip():
            break
        index += 1

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    set_core_properties(doc)
    add_masthead(doc, title, metadata)
    parse_markdown(doc, lines[index:])

    # Prevent a blank trailing paragraph from consuming a page.
    if doc.paragraphs and not doc.paragraphs[-1].text.strip():
        doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    temporary = OUTPUT.with_suffix(".part.docx")
    doc.save(temporary)
    temporary.replace(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
