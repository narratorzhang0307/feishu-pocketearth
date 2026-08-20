#!/usr/bin/env python3
"""Synchronize the authoritative Photos stage from the strategy Markdown into its DOCX peer."""
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SECTION_START = "## 阶段10A：Photos 端侧个人照片雷达"
SECTION_END = "## 阶段11：测试、性能和隐私验收"
BOARD_ROWS = ("阶段10A Photos", "阶段11 自动化验收")


def extract_section(markdown: str) -> list[str]:
    start = markdown.find(SECTION_START)
    end = markdown.find(SECTION_END, start + len(SECTION_START))
    if start < 0 or end < 0:
        raise RuntimeError("Markdown 中缺少阶段10A或阶段11边界")
    return markdown[start:end].strip().splitlines()


def extract_board_rows(markdown: str) -> dict[str, list[str]]:
    rows: dict[str, list[str]] = {}
    for line in markdown.splitlines():
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if len(cells) == 3 and cells[0] in BOARD_ROWS:
            rows[cells[0]] = cells
    missing = [name for name in BOARD_ROWS if name not in rows]
    if missing:
        raise RuntimeError(f"Markdown 执行看板缺少行：{', '.join(missing)}")
    return rows


def clean_inline(text: str) -> str:
    return text.replace("\\*", "*")


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", clean_inline(text))
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Menlo")
            run.font.size = Pt(8.5)
        else:
            paragraph.add_run(part)


def shade(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def insert_paragraph(doc: Document, before, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        add_inline(paragraph, text)
    before._p.addprevious(paragraph._p)
    return paragraph


def new_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"]._element.pPr.numPr.numId.val)
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == style_num_id)
    abstract_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    next_id = max((int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))), default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


def remove_existing_section(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    start = next((index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip().startswith("阶段10A：Photos")), None)
    if start is None:
        return
    end = next((index for index in range(start + 1, len(paragraphs)) if paragraphs[index].text.strip().startswith("阶段11：")), None)
    if end is None:
        raise RuntimeError("DOCX 中存在阶段10A，但找不到阶段11边界")
    for paragraph in paragraphs[start:end]:
        paragraph._element.getparent().remove(paragraph._element)


def sync_board_rows(doc: Document, rows: dict[str, list[str]]) -> None:
    found: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            key = row.cells[0].text.strip()
            values = rows.get(key)
            if values is None:
                continue
            if len(row.cells) != len(values):
                raise RuntimeError(f"DOCX 看板行列数不匹配：{key}")
            for cell, value in zip(row.cells, values):
                cell.text = value
            found.add(key)
    missing = [name for name in BOARD_ROWS if name not in found]
    if missing:
        raise RuntimeError(f"DOCX 执行看板缺少行：{', '.join(missing)}")


def insert_section(doc: Document, lines: list[str]) -> None:
    target = next((paragraph for paragraph in doc.paragraphs if paragraph.text.strip().startswith("阶段11：")), None)
    if target is None:
        raise RuntimeError("DOCX 中找不到阶段11插入点")
    code: list[str] = []
    in_code = False
    active_numbering: int | None = None
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                paragraph = insert_paragraph(doc, target, style="No Spacing")
                shade(paragraph, "F2F4F7")
                paragraph.paragraph_format.space_after = Pt(6)
                for index, value in enumerate(code):
                    if index:
                        paragraph.add_run().add_break(WD_BREAK.LINE)
                    run = paragraph.add_run(value)
                    run.font.name = "Menlo"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Menlo")
                    run.font.size = Pt(8)
                code = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code.append(line)
            continue
        stripped = line.strip()
        if not stripped:
            active_numbering = None
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            active_numbering = None
            level = len(heading.group(1))
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            insert_paragraph(doc, target, heading.group(2), style)
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            active_numbering = None
            insert_paragraph(doc, target, bullet.group(1), "List Bullet")
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            if active_numbering is None or numbered.group(1) == "1":
                active_numbering = new_numbering_instance(doc)
            paragraph = insert_paragraph(doc, target, numbered.group(2), "List Number")
            apply_numbering(paragraph, active_numbering)
            continue
        quote = re.match(r"^>\s?(.*)$", stripped)
        if quote:
            active_numbering = None
            insert_paragraph(doc, target, quote.group(1), "Quote")
            continue
        active_numbering = None
        insert_paragraph(doc, target, stripped, "Normal")


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("usage: sync-photos-strategy-docx.py strategy.md strategy.docx")
    md_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    markdown = md_path.read_text(encoding="utf-8")
    lines = extract_section(markdown)
    board_rows = extract_board_rows(markdown)
    doc = Document(docx_path)
    sync_board_rows(doc, board_rows)
    remove_existing_section(doc)
    insert_section(doc, lines)
    temporary = docx_path.with_suffix(".photos-sync.part.docx")
    doc.save(temporary)
    os.replace(temporary, docx_path)
    print(f"Synced {len(lines)} Markdown lines into {docx_path}")


if __name__ == "__main__":
    main()
