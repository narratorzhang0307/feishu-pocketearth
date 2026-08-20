#!/usr/bin/env python3
"""Build the hardware-Google final Word set from the confirmed round documents.

The source files stay untouched. Updated files are written into the Google
submission package so every hardware claim uses the current Gemma 4 + Gemini
implementation and carries a visual evidence reference.
"""

from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path("/Users/zhangcheng/Desktop/pocket earth_google")
SOURCE = Path("/Users/zhangcheng/Desktop/未命名文件夹 12")
OUT = ROOT / "PocketEarthGoogle_提交包" / "07_统一审核材料_硬件Google更新版"
BOARD = ROOT / "docs/assets/hardware/core-4k/05_Gemma与Gemini端云双脑_4K.png"
BOARD_DEVICE = ROOT / "docs/assets/hardware/core-4k/03_三个真实硬件入口_4K.png"


def set_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def set_cell(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)


def add_figure_before(document: Document, anchor, image: Path, caption: str) -> None:
    figure = document.add_paragraph()
    figure.alignment = 1
    figure.add_run().add_picture(str(image), width=Inches(6.25))
    anchor._p.addprevious(figure._p)
    cap = document.add_paragraph()
    cap.alignment = 1
    run = cap.add_run(caption)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.italic = True
    anchor._p.addprevious(cap._p)


def save(document: Document, filename: str) -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    target = OUT / filename
    document.save(target)
    return target


def update_c() -> Path:
    src = SOURCE / "Pocket_Earth_C轮_解决方案与可运行Demo_最终确认版.docx"
    doc = Document(src)
    set_cell(
        doc.tables[1].rows[4].cells[0],
        "5. Frost Edge 软硬件共生。用户操作：在 Raspberry Pi 5 × Whisplay 实体设备中选择口袋播客、日落电台或地球答案，也可打开单文件数字孪生复核十二张真实设备界面。系统处理：Google Gemma 4 E4B IT QAT Q4_0 通过 127.0.0.1:8787 回环服务承担本地分类、受限选择、短回复与弱网降级；复杂且符合公共事件边界的任务再升级 Google Gemini。设备只同步白名单公共事件和可缓存内容，不保存私人原文、原图、完整画像、精确坐标或云密钥。最终结果与核心价值：Frost 在网页和实体端保持同一人格与公共知识版次，硬件具备可核验的本地 Google AI 推理，同时与私人地球维持清晰边界。",
    )
    set_text(
        doc.paragraphs[33],
        "Frost Edge 在线状态：硬件整机、树莓派端软件、Gemma 4 E4B 独立回环服务与三入口 Launcher 均已有代码和设备材料。当前线上 Frost Feed 关闭，公开 Demo 与实体设备未保持实时连接；最终树莓派真实推理验收仍待模型传输完成后补充 GEMMA-4-E4B-VALIDATION.md，现有材料不将其写成已验证。",
    )
    set_text(
        doc.paragraphs[18],
        doc.paragraphs[18].text + " Frost Edge 在线数字孪生：https://pocketearth-google.throughtheglass.art/hardware-digital-twin.html。",
    )
    add_figure_before(doc, doc.paragraphs[34], BOARD_DEVICE, "图 C-1　Frost Edge 三个真实硬件入口：口袋播客、日落电台、地球答案")
    return save(doc, "Pocket_Earth_C轮_解决方案与可运行Demo_硬件Google更新版.docx")


def update_d() -> Path:
    src = SOURCE / "Pocket_Earth_D轮_Google技术应用_最终提交版.docx"
    doc = Document(src)
    set_text(
        doc.paragraphs[5],
        "当前可据代码路径、部署记录、模型校验值与产品文档确认五组 Google 技术：浏览器端 Gemma 3n E2B IT、MediaPipe Tasks GenAI / LlmInference、树莓派端 Gemma 4 E4B IT QAT Q4_0、Gemini 云端模型族，以及 Google Gemini API / Google AI Studio Key 官方接入路径。Harness、RunTrace、Validator、Confirm Gate 与 Frost Edge 的产品编排由项目自研；Mapbox、WebGPU 与备用传输属于外部基础设施，不计入 Google 技术清单。",
    )
    set_text(doc.paragraphs[34], "Frost Edge 在 Raspberry Pi 5 本机运行 Google Gemma 4 E4B IT QAT Q4_0。模型文件 gemma-4-E4B_q4_0-it.gguf 为 5,154,941,280 字节，SHA-256 为 676c35070db6dbe52f93e9c864ee0fba4eddea94b9c875d9cb10daff453fbaee；独立 pocket-earth-gemma.service 仅绑定 127.0.0.1:8787/v1。设备 Harness 用它完成本地分类、公共低风险选择、短回复和弱网降级，复杂且经边界允许的任务才升级 Gemini。云密钥不会进入设备界面、仓库或公共事件。")
    set_text(doc.paragraphs[41], "当前未使用、不能在审核中声称使用的 Google 技术包括 Google Search Grounding、Vertex AI、Firebase、Flutter、Android Studio、Google ADK、端侧向量搜索和端侧 RAG。WebGPU 属于浏览器标准；Mapbox、OpenStreetMap、Open-Meteo、Unsplash、KIRI 与备用传输属于非 Google 技术。项目的 Google 技术深度由浏览器 Gemma 3n、MediaPipe、树莓派 Gemma 4 与云端 Gemini 的真实分工共同构成。")
    set_text(doc.paragraphs[42], "截至 2026-07-19 的实验、维护与动态状态：LiteRT-LM JavaScript 迁移仍为预留路线；Gemini 图像生成和 3D / AR 属于可选增强；Google 官方 Gemini API 路径已实现，当前线上可使用 Google-only 备用传输；Frost Edge 整机、Google 版设备代码、Gemma 4 独立服务、十二张设备界面与数字孪生均已形成材料，但最终树莓派真实推理验收仍待 GEMMA-4-E4B-VALIDATION.md；MediaPipe 锁定为 0.10.29。线上状态以 /healthz 为准。")

    table = doc.tables[1]
    row = table.add_row()
    set_cell(row.cells[0], "3. Frost Edge 端侧模型")
    set_cell(row.cells[1], "Google Gemma 4 E4B IT QAT Q4_0；Raspberry Pi 5 本地 GGUF 权重，独立 systemd 服务，127.0.0.1 回环 API。")
    table.rows[3]._tr.addprevious(row._tr)
    # Renumber the cloud/API rows after the inserted hardware model row.
    set_cell(table.rows[4].cells[0], "4. Gemini 云端模型族")
    set_cell(table.rows[5].cells[0], "5. Gemini 官方接入路径")

    set_cell(doc.tables[2].rows[4].cells[1], "浏览器主流程使用 Gemma 3n + MediaPipe；Frost Edge 使用单独的 Gemma 4 E4B 回环服务。两端都只提出候选结果，私人写入、公共发布和硬件动作分别由 FactRelay、Confirm Gate 与设备白名单负责。")
    set_cell(doc.tables[9].rows[6].cells[1], "hardware/frost-edge-google/raspi/：Gemma 4 安装器、systemd、回环客户端、三入口 Launcher、事件桥、preflight 与 smoke test；hardware/frost-edge-google/sunset-radio/：完整设备源码快照；docs/assets/hardware/ 与 public/hardware-digital-twin.html：视觉及交互证据。")
    set_cell(doc.tables[10].add_row().cells[0], "树莓派 Gemma 4 服务未就绪")
    row = doc.tables[10].rows[-1]
    set_cell(row.cells[1], "三种硬件入口继续使用确定性规则、本地目录和上一有效公共缓存；设备 preflight 标记 local-gemma unavailable。")
    set_cell(row.cells[2], "不把规则或缓存伪装成 Gemma 结果，不因本地失败静默上传私人数据，也不阻断固定硬件入口。")
    set_cell(doc.tables[11].rows[0].cells[0], "D 轮技术结论：Pocket Earth 的 Google 技术主链由浏览器 Gemma 3n E2B IT、MediaPipe Tasks GenAI、树莓派 Gemma 4 E4B IT、Gemini 按任务分层模型与 Gemini 官方接入路径组成。Frost Agent Harness 决定在哪一层执行，Boundary 和 Confirm Gate约束上传、写入与硬件动作，RunTrace 与设备 preflight 提供可审计证据。")
    add_figure_before(doc, doc.paragraphs[41], BOARD, "图 D-1　Google Gemma 与 Gemini 在浏览器、树莓派和云端的分工")
    return save(doc, "Pocket_Earth_D轮_Google技术应用_硬件Google更新版.docx")


def update_e() -> Path:
    src = SOURCE / "Pocket_Earth_E轮_创新与海外本地化_最终确认版.docx"
    doc = Document(src)
    set_cell(
        doc.tables[1].rows[4].cells[1],
        "常见 AI 硬件演示依赖远端接口或网页镜像，设备本身缺少可核验的模型服务。Frost Edge 在 Raspberry Pi 5 上运行 Google Gemma 4 E4B 独立回环服务，本地完成分类、受限选择、短回复和弱网降级；复杂公共任务经边界检查后升级 Gemini。口袋播客、日落电台和地球答案共享 Frost 人格、公共知识版次、事件白名单与缓存，设备拒绝私人原文、原图、完整画像、精确坐标和云密钥。用户获得可触摸的 Google AI 体验，私人知识仍留在软件端。证据：hardware/frost-edge-google/、docs/hardware/FROST-EDGE-GOOGLE.md、docs/assets/hardware/ 与硬件数字孪生。",
    )
    set_text(doc.paragraphs[42], "• 产品实现证据：公开 PWA、看展主闭环、公私双地球、Gemma/Gemini Harness、RunTrace、公共知识与 Frost Edge 均有代码或运行材料；当前 Web 基线为 52 个测试文件、1336 项测试通过，硬件另有 Gemma 适配器、三入口、事件白名单、smoke test、设备 preflight、十二张 Whisplay 界面、七张提交证据图与单文件数字孪生。")
    add_figure_before(doc, doc.paragraphs[42], BOARD_DEVICE, "图 E-1　实体端的三个公共、低风险、可缓存入口")
    return save(doc, "Pocket_Earth_E轮_创新与海外本地化_硬件Google更新版.docx")


def update_f() -> Path:
    src = SOURCE / "Pocket_Earth_F轮_视频与声明_最终确认版.docx"
    doc = Document(src)
    set_cell(doc.tables[4].rows[9].cells[1], "完成 Frost Edge 整机、树莓派三入口 Launcher、Whisplay UI、驱动、公共知识同步、事件桥和 smoke test；新增 Google Gemma 4 E4B IT QAT Q4_0 独立回环服务、设备端适配器、安装与 SHA-256 校验、systemd 单元、preflight 和 Gemini 受控升级路径。形成日落电台、口袋播客、地球答案三个根入口，并将设备事件限制为公共、低风险、可缓存白名单。当前线上 Frost Feed 关闭，提交材料不把硬件描述为与公开 Demo 实时连接。")
    set_cell(doc.tables[4].rows[10].cells[1], "将 Google 版部署到独立域名、目录、端口、PM2 进程与知识 Worker；完成 HTTPS、/healthz、浏览器 Gemma Range、provider/model 状态和公开访问验证。完成 TypeScript 检查、生产构建、公共知识验证及 52 个测试文件、1336 项 Web 测试；补充 Frost Edge 真机代码、硬件 smoke test、模型端点核验、十二张设备界面、七张硬件证据图、在线与离线数字孪生。")
    note = doc.add_paragraph()
    note.add_run("硬件补充证据说明：").bold = True
    note.add_run("已发布的 07:47 Demo 保持原时间戳，因此视频时间码不追加硬件画面。Frost Edge 通过 GitHub hardware/frost-edge-google/、docs/assets/hardware/、硬件技术说明、真机屏幕和数字孪生补充核验。在线入口：https://pocketearth-google.throughtheglass.art/hardware-digital-twin.html。")
    doc.tables[2]._tbl.addnext(note._p)
    add_figure_before(doc, doc.paragraphs[14], BOARD, "图 F-1　Frost Edge 的 Google 端云双脑证据")
    return save(doc, "Pocket_Earth_F轮_视频与声明_硬件Google更新版.docx")


def update_product_doc() -> Path:
    src = ROOT / "PocketEarthGoogle_提交包/06_产品文档/Pocket_Earth_口袋地球_完整产品文档_Google参赛版.docx"
    doc = Document(src)
    set_text(doc.paragraphs[8], "联系信息：未纳入公开仓库")
    set_text(doc.paragraphs[220], "Frost Edge 是 Pocket Earth 的实体陪伴端，将公共、低风险、可缓存的空间知识能力带到 Raspberry Pi 5 × Whisplay。设备保留自己的 Google Gemma 推理平面、固定入口、目录与缓存；Private Earth 的个人原文、原图、完整画像、精确坐标和云密钥不进入硬件事件。")
    set_text(doc.paragraphs[223], "10.3 Google Gemma × Gemini 硬件双脑")
    set_text(doc.paragraphs[224], "Google Gemma 4 E4B IT QAT Q4_0 在树莓派本机运行。模型文件为 gemma-4-E4B_q4_0-it.gguf，大小 5,154,941,280 字节，SHA-256 为 676c35070db6dbe52f93e9c864ee0fba4eddea94b9c875d9cb10daff453fbaee。pocket-earth-gemma.service 仅绑定 127.0.0.1:8787/v1，负责本地分类、受限选择、短回复与弱网降级；复杂且符合公共事件边界的任务才升级到 Google Gemini。设备 Harness 保存 provider、modelOwner、transport、耗时与 fallback。")
    set_text(doc.paragraphs[225], "10.4 三入口、统一人格与事件安全")
    set_text(doc.paragraphs[226], "口袋播客、日落电台和地球答案共用 Frost 人格、Whisplay UI、公共知识版次、事件白名单与本地缓存。设备桥只允许 music_now_playing、public_knowledge_brief 和 buddy_status 等公共事件，并拒绝凭证形态内容。当前线上 Frost Feed 未开启，公开 Demo 与实体设备未保持实时连接。")
    set_text(doc.paragraphs[227], "10.5 当前完成状态与核验证据")
    set_text(doc.paragraphs[228], "硬件整机、树莓派端软件、Google Gemma 4 独立服务、三入口 Launcher、Gemini 受控升级、smoke test、真实 Whisplay 界面与数字孪生均有材料。最终树莓派真实推理验收仍待模型传输完成后补充 hardware/frost-edge-google/raspi/GEMMA-4-E4B-VALIDATION.md；在该报告出现前不宣称已经完成真机最终验证。")
    add_figure_before(doc, doc.paragraphs[230], BOARD, "图 10-1　Frost Edge：Google Gemma 4 本地推理与 Gemini 按需升级")
    add_figure_before(doc, doc.paragraphs[230], BOARD_DEVICE, "图 10-2　Frost Edge 三个真实入口与 Whisplay 状态")
    return save(doc, "Pocket_Earth_口袋地球_完整产品文档_Google硬件深化版.docx")


def main() -> None:
    outputs = [update_c(), update_d(), update_e(), update_f(), update_product_doc()]
    copy2(ROOT / "树莓派硬件/Pocket_Earth_Frost_Edge_数字孪生_单文件版.html", OUT / "Pocket_Earth_Frost_Edge_数字孪生_单文件版.html")
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
