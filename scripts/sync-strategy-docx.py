#!/usr/bin/env python3
"""Insert the 2026-08-11 execution dashboard into the existing strategy DOCX."""

from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs/strategy/Pocket Earth 决赛改造总计划与执行准则.docx"

ROWS = [
    ("P0、阶段9 书影音", "已完成并冻结", "书籍、电影、音乐均按 pocket-data/v1 安装、切换、落位和卸载；本轮只做回归。"),
    ("阶段0 基线冻结", "已完成", "基线提交 848ba12；为保护并行修改，未强行切分支或打标签。"),
    ("阶段1 命名与信息架构", "已完成", "一个 Frost Agent + Skills；公共知识层退出活跃页面。"),
    ("阶段2 协议与 Registry", "已完成", "pocket-skill/v1 严格校验、Registry、生命周期、示例和测试已落地。"),
    ("阶段3 OSS 与首屏", "部分完成 / 上传受阻", "资产清单、Loader、进度/取消/续传、视野聚合和首屏门禁完成；实际上传受本机 STS 过期阻塞。"),
    ("阶段4 云端 Qwen", "已完成", "生产/Vite 共用 DashScope Qwen Provider；旧接口 410，Gemma Web runtime/依赖已移除。"),
    ("阶段5 Android Qwen/MNN", "部分完成 / 真机受阻", "Java/JNI/Capacitor 契约和 Android 36 debug APK 编译通过；缺签名 JNI、目标 Armv9 手机与 SME2 A/B。"),
    ("阶段6 Plaza / RunTrace", "已完成", "安装状态、权限、质量门禁、资产生命周期和结构化 RunTrace 已接通。"),
    ("阶段7 专业 Skills", "开发机范围完成", "用户指定冻结旅行/看展；碑拓识读与复原完成真实 MNN/LoRA/Quality Gate 闭环。"),
    ("阶段8 内容 Mapping", "已完成并冻结", "Book-to-Earth 已完成，本轮不重做。"),
    ("阶段10 UI 适配", "已完成本轮范围", "Skills、Plaza、模型管理、碑拓与地图保持 Pocket Earth 视觉；移动端浏览器闭环通过。"),
    ("阶段10A Photos", "并行任务负责", "本轮不修改 Photos Tab 与照片推理，避免覆盖另一 Codex 窗口。"),
    ("阶段11 自动化验收", "Web/Android build 完成", "77 files / 1,451 tests；首屏 836,300 bytes；APK 30,441,299 bytes；真机推理待补。"),
    ("阶段12 决赛证据", "已建立，真机项待补", "实施状态与 SME2 A/B 协议已写入 docs/evidence。"),
    ("阶段13 清理发布", "部分完成", "活跃旧模型运行时和移动包重资产已清理；Release 签名、OSS 实传和真机记录待补。"),
]


def set_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table) -> None:
    widths = [2050, 1900, 5410]
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for cell in header.cells:
        set_shading(cell, "111111")
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "Arial"
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
                run.font.size = Pt(8.5)
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if r_idx > 0:
            status = row.cells[1].text
            set_shading(row.cells[1], "E7F4EC" if status.startswith("已完成") or "完成" in status and "部分" not in status else "FFF1C9")


def main() -> None:
    doc = Document(DOCX)
    if any(p.text.strip() == "0.1 执行看板（2026-08-11）" for p in doc.paragraphs):
        print("strategy dashboard already present")
        return
    anchor = next(p for p in doc.paragraphs if p.text.strip().startswith("Word 文件是阶段性快照"))
    heading = doc.add_paragraph("0.1 执行看板（2026-08-11）", style="Heading 2")
    anchor._p.addnext(heading._p)
    table = doc.add_table(rows=1, cols=3)
    for cell, text in zip(table.rows[0].cells, ("阶段", "状态", "本轮结果 / 证据")):
        cell.text = text
    for values in ROWS:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(table)
    heading._p.addnext(table._tbl)
    note = doc.add_paragraph()
    run = note.add_run("统一证据索引：docs/evidence/implementation-status-20260811.md。Android JNI、飞行模式、OSS 实传或 SME2 不得由桌面结果推断。")
    run.italic = True
    run.font.size = Pt(9)
    table._tbl.addnext(note._p)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
